#!/usr/bin/env python3
"""
江苏科技大学本科毕业设计（论文）格式自动排版工具
基于《毕业设计(论文)撰写规范(理工类)》实现
"""
import argparse
import re
import os
import sys

# 解决 Windows 终端 GBK 编码问题
if sys.stdout.encoding and sys.stdout.encoding.upper() in ('GBK', 'CP936'):
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)


# ============ 默认配置 ============
DEFAULT_CONFIG = {
    "school": "江苏科技大学",
    "top_margin_cm": 2.54,
    "bottom_margin_cm": 2.54,
    "left_margin_cm": 3.0,
    "right_margin_cm": 2.0,
    "binding_cm": 0.5,
    "line_spacing": 1.5,
}

# ============ 字号映射（Pt） ============
FS = {
    "小二": Pt(18), "小二号": Pt(18),
    "小三": Pt(15), "小三号": Pt(15),
    "小四": Pt(12), "小四号": Pt(12),
    "四号": Pt(14), "小五号": Pt(9),
    "五号": Pt(10.5), "三号": Pt(16), "二号": Pt(22),
}


# ============ 底层工具函数 ============

def _set_run_font(run, font_cn="宋体", font_en="Times New Roman", size=None, bold=None):
    """设置单个 run 的字体"""
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def _set_para_spacing(para, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落间距"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)


def _set_para_indent(para, first_line_chars=0, font_size_pt=12):
    """设置首行缩进"""
    para.paragraph_format.first_line_indent = Pt(font_size_pt * first_line_chars) if first_line_chars else Pt(0)


def _set_para_font(para, font_cn="宋体", font_en="Times New Roman", size=None, bold=None):
    """设置段落内所有 run 的字体"""
    for run in para.runs:
        _set_run_font(run, font_cn, font_en, size, bold)


# ============ 页面设置 ============

def setup_page(doc, config):
    """设置页面布局"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(config["top_margin_cm"])
    section.bottom_margin = Cm(config["bottom_margin_cm"])
    section.left_margin = Cm(config["left_margin_cm"] + config["binding_cm"])
    section.right_margin = Cm(config["right_margin_cm"])


# ============ 样式修改 + 段落级覆盖 ============

def modify_styles_and_format_paragraphs(doc):
    """修改 Word 内置样式 + 段落级格式覆盖（双保险）"""
    styles = doc.styles

    # ------- 修改样式 -------
    for style_name, font_cn, font_en, size, bold, aligns, ls, sb, sa, indent in [
        ("Normal",     "宋体", "Times New Roman", FS["小四"],   None,
         WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, 0, 0, 480),
        ("Heading 1",  "宋体", "Times New Roman", FS["小二号"], True,
         WD_ALIGN_PARAGRAPH.CENTER,  1.5, 9, 9, 0),
        ("Heading 2",  "宋体", "Times New Roman", FS["小三"],   True,
         WD_ALIGN_PARAGRAPH.LEFT,    1.5, 0, 7, 0),
        ("Heading 3",  "宋体", "Times New Roman", FS["小三"],   True,
         WD_ALIGN_PARAGRAPH.LEFT,    1.5, 0, 7, 0),
        ("Caption",    "宋体", "Times New Roman", FS["五号"],   None,
         WD_ALIGN_PARAGRAPH.CENTER,  1.5, 0, 0, 0),
    ]:
        if style_name not in styles:
            continue
        s = styles[style_name]
        # 字体
        s.font.name = font_en
        s.font.size = size
        if bold is not None:
            s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        rPr = s.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            s.element.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)

        # 段落属性（通过 XML 直接写入 pPr）
        pPr = s.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            s.element.insert(0, pPr)

        # 对齐
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = parse_xml(f'<w:jc {nsdecls("w")}></w:jc>')
            pPr.append(jc)
        align_map = {WD_ALIGN_PARAGRAPH.CENTER: "center", WD_ALIGN_PARAGRAPH.LEFT: "left",
                     WD_ALIGN_PARAGRAPH.JUSTIFY: "both"}
        jc.set(qn('w:val'), align_map.get(aligns, "left"))

        # 间距
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}></w:spacing>')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(int(ls * 240)))
        spacing.set(qn('w:lineRule'), "auto")
        if sb is not None:
            spacing.set(qn('w:before'), str(sb * 20))  # pt→twips
        if sa is not None:
            spacing.set(qn('w:after'), str(sa * 20))

        # 缩进
        if indent is not None and indent > 0:
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = parse_xml(f'<w:ind {nsdecls("w")}></w:ind>')
                pPr.append(ind)
            ind.set(qn('w:firstLine'), str(indent))
        elif indent == 0:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)

    # ------- 段落级覆盖 -------
    for para in doc.paragraphs:
        text = para.text.strip()
        sn = para.style.name

        # 文本内容检测: 无编号章节标题
        if re.match(r'^第\d+章$', text) and sn != "Heading 1":
            para.style = styles["Heading 1"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_font(para, "宋体", "Times New Roman", FS["小二号"], bold=True)
            _set_para_spacing(para, 1.5, 9, 9)
            _set_para_indent(para, 0)
        elif text == "本章小结" and sn in ("Normal", "MTDisplayEquation"):
            para.style = styles["Heading 2"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_font(para, "宋体", "Times New Roman", FS["小三"], bold=True)
            _set_para_spacing(para, 1.5, 0, 7)
            _set_para_indent(para, 0)
        elif text in ("参考文献", "参 考 文 献") and sn not in ("Heading 1",):
            para.style = styles["Heading 1"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_font(para, "宋体", "Times New Roman", FS["小二号"], bold=True)
            _set_para_spacing(para, 1.5, 9, 9)
            _set_para_indent(para, 0)
        elif sn == "Heading 1":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_font(para, "宋体", "Times New Roman", FS["小二号"], bold=True)
            _set_para_spacing(para, 1.5, 9, 9)
            _set_para_indent(para, 0)
        elif sn == "Heading 2":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_font(para, "宋体", "Times New Roman", FS["小三"], bold=True)
            _set_para_spacing(para, 1.5, 0, 7)
            _set_para_indent(para, 0)
        elif sn == "Heading 3":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_font(para, "宋体", "Times New Roman", FS["小三"], bold=True)
            _set_para_spacing(para, 1.5, 0, 7)
            _set_para_indent(para, 0)
        elif sn == "Caption":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_font(para, "宋体", "Times New Roman", FS["五号"])
            _set_para_spacing(para, 1.5)
            _set_para_indent(para, 0)
        elif sn == "Normal" and text and not text.startswith("关键词") and not text.startswith("Keywords"):
            # 跳过封面/摘要区域的段落（不覆盖它们的显式格式）
            is_cover_or_abstract = any(kw in text for kw in ["江苏科技大学", "本 科 毕 业 设 计"])
            is_abstract_title = text.replace(" ", "") in ("摘要", "Abstract")
            is_abstract_body = text.endswith("。") and len(text) > 50 and not text.startswith("图")
            if is_cover_or_abstract or is_abstract_title:
                pass  # 保留已有的显式格式
            elif is_abstract_body:
                pass  # 保留已有的显式格式
            else:
                # 正文段落：清除显式格式，让其继承 Normal 样式
                for run in para.runs:
                    run.font.size = None
                    run.font.bold = None
                    run.font.name = None
                    try:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    except:
                        pass
                    r = run._element
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            rFonts.set(qn('w:eastAsia'), "宋体")
                para.paragraph_format.line_spacing = 1.5
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ============ 特殊段落格式化 ============

def format_abstract(doc):
    """格式化摘要标题与关键词"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.replace(" ", "") in ("摘要", "Abstract"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_font(para, "宋体", "Times New Roman", FS["小二"], bold=True)
            _set_para_spacing(para, 1.5, 6, 6)
            _set_para_indent(para, 0)
        elif text.startswith("关键词"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_font(para, "宋体", "Times New Roman", FS["三号"], bold=True)
            _set_para_spacing(para, 1.5)
            _set_para_indent(para, 0)
        elif text.startswith("Keywords"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_font(para, "宋体", "Times New Roman", FS["三号"], bold=True)
            _set_para_spacing(para, 1.5)
            _set_para_indent(para, 0)


def format_references(doc):
    """格式化参考文献条目（5号宋体）"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'^\[\d+\]', text):
            _set_para_font(para, "宋体", "Times New Roman", FS["五号"])
            _set_para_spacing(para, 1.5)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_indent(para, 0)


def fix_cover_fields(doc):
    """封面字段 → 四号黑体"""
    for para in doc.paragraphs:
        text = para.text.strip()
        for kw in ("学    院", "专    业", "学生姓名", "班级学号", "指导教师"):
            if text.startswith(kw):
                _set_para_font(para, "黑体", "Times New Roman", FS["四号"])
                break


def format_tables(doc):
    """格式化表格为三线表"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_para_font(para, "宋体", "Times New Roman", FS["五号"])
                    para.paragraph_format.line_spacing = 1.0
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if len(table.rows) == 0:
            continue

        # 表头上下边框
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'
            )
            old = tcPr.find(qn('w:tcBorders'))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(borders)

        # 表末下边框
        for cell in table.rows[-1].cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'
            )
            old = tcPr.find(qn('w:tcBorders'))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(borders)


# ============ 主流程 ============

def format_thesis(input_path, output_path=None, config=None):
    if config is None:
        config = DEFAULT_CONFIG

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_已排版{ext}"

    doc = Document(input_path)
    print(f"输入: {input_path}")
    print(f"输出: {output_path}\n")

    print("[1/7] 页面设置 (A4, 上下2.54cm, 左3.0cm, 右2.0cm)...")
    setup_page(doc, config)

    print("[2/7] 修改样式 + 段落级格式覆盖...")
    modify_styles_and_format_paragraphs(doc)

    print("[3/7] 格式化摘要与关键词...")
    format_abstract(doc)

    print("[4/7] 格式化封面字段...")
    fix_cover_fields(doc)

    print("[5/7] 格式化表格 (三线表)...")
    format_tables(doc)

    print("[6/7] 格式化参考文献...")
    format_references(doc)

    print("[7/7] 保存文件...")
    doc.save(output_path)
    print(f"\n完成! 文件已保存: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="江苏科技大学本科毕业设计(论文)格式自动排版工具"
    )
    parser.add_argument("input", help="输入 docx 文件路径")
    parser.add_argument("-o", "--output", help="输出文件路径")
    parser.add_argument("--school", default="江苏科技大学")
    parser.add_argument("--top", type=float, default=2.54, help="上边距 (cm)")
    parser.add_argument("--bottom", type=float, default=2.54, help="下边距 (cm)")
    parser.add_argument("--left", type=float, default=3.0, help="左边距 (cm)")
    parser.add_argument("--right", type=float, default=2.0, help="右边距 (cm)")
    parser.add_argument("--binding", type=float, default=0.5, help="装订线 (cm)")
    parser.add_argument("--line-spacing", type=float, default=1.5, help="行距倍数")

    args = parser.parse_args()
    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)

    config = {k: getattr(args, k) for k in ["school", "line_spacing"]}
    config.update({
        "top_margin_cm": args.top,
        "bottom_margin_cm": args.bottom,
        "left_margin_cm": args.left,
        "right_margin_cm": args.right,
        "binding_cm": args.binding,
    })
    format_thesis(args.input, args.output, config)


if __name__ == "__main__":
    main()
